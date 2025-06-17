print("📂 LOADED transcribe_video.py from:", __file__)

# C:\Users\X\PycharmProjects\video-translator-gibuy-main\backend\transcribe_video.py
import whisper
import os
import subprocess
import logging
from datetime import timedelta
from pathlib import Path
import json
import sys
import openai
from docx import Document
from fpdf import FPDF
import cv2  # שים למעלה אם עדיין לא מיובא
from r2_client import upload_file_to_r2
from billing import is_allowed_to_transcribe, PlanType
from mongo_client import get_user_usage_and_plan, update_user_usage
from datetime import datetime
from mongo_client import transcriptions_collection

import re
from textwrap import wrap
from utils import translate_segments

def convert_srt_to_ass(srt_path, ass_path, resolution=(1280, 720), style_config=None, rtl=False):
    import os
    import pysubs2

    subs = pysubs2.load(srt_path, encoding="utf-8")
    subs.video_width, subs.video_height = resolution
    style = subs.styles["Default"]

    style.fontname = style_config.get("font", "Noto Sans")
    style.fontsize = style_config.get("size", 36)

    alignment_map = {
        "bottom_center": 2,
        "bottom_right": 3,
        "bottom_left": 1
    }

    style.alignment = alignment_map.get(style_config.get("alignment", "bottom_center"), 2)
    style.margin_v = style_config.get("margin_v", 60)
    style.outline = style_config.get("outline", 2)
    style.shadow = style_config.get("shadow", 0)

    def hex_to_color(hx):
        hx = hx.lstrip("#")
        return pysubs2.Color(int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16))

    style.primarycolor = hex_to_color(style_config.get("primary_color", "#FFFFFF"))
    style.outlinecolor = hex_to_color(style_config.get("outline_color", "#000000"))

    if rtl:
        style.alignment = alignment_map.get("bottom_right", 3)
        style.direction = "rtl"

    # שמירת קובץ ASS
    # שמור גם עותק נוסף בתיקיית results:
    permanent_ass_path = os.path.join("results", "debug_subs.ass")
    subs.save(ass_path)
    subs.save(permanent_ass_path)
    print(f"📁 עותק נוסף של ASS נשמר גם כאן: {permanent_ass_path}")

    # בדיקת קיום הקובץ
    if not os.path.exists(ass_path):
        raise FileNotFoundError(f"קובץ .ass לא נוצר: {ass_path}")

    # בדיקה שהתוכן לא ריק
    if os.path.getsize(ass_path) == 0:
        raise ValueError(f"קובץ .ass קיים אך ריק: {ass_path}")

    print(f"✅ קובץ ASS נוצר בהצלחה: {ass_path}")



import cv2

import subprocess
import os

def burn_ass_subtitles(video_path, ass_path, output_path, resolution=None):
    """
    מטמיע כתוביות מסוג ASS לתוך וידאו באמצעות FFmpeg.

    :param video_path: מיקום קובץ הווידאו (str)
    :param ass_path: מיקום קובץ ה־.ass (str)
    :param output_path: מיקום לקובץ הווידאו המוגמר עם כתוביות (str)
    :param resolution: (לא בשימוש כרגע, אבל אפשר להרחיב בעתיד)
    """

    # המרה לסלשים קדמיים ל־FFmpeg (חיוני ב-Windows)
    ass_path_fixed = ass_path.replace("\\", "/")

    # מסנן וידאו של FFmpeg ל־ASS
    vf_filter = f"ass={ass_path_fixed}"

    # פקודת FFmpeg
    cmd = [
        "ffmpeg", "-y",
        "-i", video_path,
        "-vf", vf_filter,
        "-c:a", "copy",
        output_path
    ]

    print("Running command:", " ".join(cmd))

    # הרצת FFmpeg (לא דרך shell!)
    subprocess.run(cmd, check=True)


# הגדרת PATH ל־ffmpeg
os.environ["PATH"] = r"C:\\ffmpeg-2025-05-15-git-12b853530a-full_build\\bin" + os.pathsep + os.environ["PATH"]

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)
logging.getLogger("celery").setLevel(logging.ERROR)
logging.getLogger("kombu").setLevel(logging.ERROR)


def update_status(step, progress):
    with open("../status.json", "w", encoding="utf-8") as f:
        json.dump({"step": step, "progress": progress}, f)

class Config:
    def __init__(self):
        self.load_config()

    def load_config(self):
        default_config = {
            "openai_api_key": "",
            "whisper_model": "large-v3",
            "ffmpeg_path": "C:/ffmpeg-2025-05-15-git-12b853530a-full_build/bin/ffmpeg.exe",
            "input_video": "video.mp4",
            "output_audio": "temp_audio.wav",
            "output_srt": "subtitles.srt",
            "language": "auto",
            "translation_target": None,
            "compute_type": "float16",
            "output_mode": "burn",
            "subtitle_style": {
                "font": "Noto Sans",
                "size": 36,
                "alignment": "bottom_center",
                "primary_color": "#FFFFFF",
                "outline_color": "#000000",
                "outline": 2,
                "shadow": 0,
                "margin_v": 60
            }
        }

        try:
            with open('config.json', 'r', encoding='utf-8') as f:
                user_config = json.load(f)

            # מיזוג חכם: מה שב־config.json גובר, השאר מברירת מחדל
            merged_config = default_config.copy()
            merged_config.update(user_config)
            self.__dict__.update(merged_config)

        except FileNotFoundError:
            logger.warning("⚠️ config.json לא נמצא, משתמש בברירת מחדל")
            self.__dict__.update(default_config)

        except json.JSONDecodeError:
            logger.error("❌ שגיאה בפורמט config.json")
            sys.exit(1)


class HebrewTranscriber:
    def __init__(self, config):
        self.config = config
        self.setup_ffmpeg()
        self.FFPROBE_PATH = r"C:\ffmpeg-2025-05-15-git-12b853530a-full_build\bin\ffprobe.exe"


    def setup_ffmpeg(self):
        try:
            subprocess.run([self.config.ffmpeg_path, "-version"], capture_output=True, check=True)
            logger.info("✅ FFmpeg זוהה ומוכן")
        except (subprocess.CalledProcessError, FileNotFoundError):
            logger.error("❌ FFmpeg לא נמצא. ודא שהנתיב נכון")
            sys.exit(1)




    def extract_audio(self):
        logger.info(f"🎵 חילוץ ושיפור אודיו מ־{self.config.input_video}")
        try:
            subprocess.run([
                self.config.ffmpeg_path,
                "-y",
                "-i", self.config.input_video,
                "-vn",
                "-af", "highpass=f=200, lowpass=f=3000, afftdn=nf=-25",  # מסנן רעשים
                "-acodec", "pcm_s16le",
                "-ar", "16000",
                "-ac", "1",
                self.config.output_audio
            ], check=True, capture_output=True, text=True)
            logger.info(f"✅ אודיו נוצר: {self.config.output_audio}")
        except Exception as e:
            logger.error(f"❌ שגיאה בחילוץ האודיו: {str(e)}")
            raise

    def transcribe_audio(self):
        logger.info("🧠 הרצת תמלול עם Whisper")
        try:
            kwargs = {}
            if self.config.whisper_model.startswith("faster"):
                kwargs["compute_type"] = getattr(self.config, "compute_type", "float16")

            model = whisper.load_model(
                self.config.whisper_model,
                device="cuda",
                **kwargs
            )

            task_mode = "translate" if self.config.translation_target else "transcribe"
            language = None if self.config.language == "auto" else self.config.language

            result = model.transcribe(
                self.config.output_audio,
                language=language,
                word_timestamps=True,
                task=task_mode
            )

            words = result.get("words", [])
            segments = []

            if words:
                max_segment_duration = 6.0
                max_wps = 3.0  # מילים לשנייה = 180 מילים בדקה
                gap_threshold = 1.2

                current = []
                for i, word in enumerate(words):
                    current.append(word)
                    is_last = i == len(words) - 1
                    gap = (
                        words[i + 1]["start"] - word["end"]
                        if not is_last
                        else 0
                    )
                    duration = current[-1]["end"] - current[0]["start"]
                    wps = len(current) / duration if duration > 0 else 0

                    should_split = (
                            gap > gap_threshold or
                            duration >= max_segment_duration or
                            wps > max_wps or
                            is_last
                    )

                    if should_split:
                        segments.append({
                            "start": current[0]["start"],
                            "end": current[-1]["end"],
                            "text": " ".join([w["word"] for w in current])
                        })
                        current = []

            else:
                segments = result["segments"]

            self.source_language = result.get("language", "unknown")
            return segments

        except Exception as e:
            logger.error(f"❌ שגיאה בתמלול: {str(e)}")
            raise

    def format_timestamp(self, seconds):
        td = timedelta(seconds=seconds)
        total = str(td)
        if '.' in total:
            total = total.split('.')[0] + ',' + total.split('.')[1][:3]
        else:
            total = total + ',000'
        if len(total.split(':')[0]) == 1:
            total = "0" + total
        return total

    def generate_txt(self, segments, output_path):
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                for seg in segments:
                    f.write(seg["text"].strip() + "\n")
            logger.info(f"✅ קובץ TXT נוצר: {output_path}")
        except Exception as e:
            logger.error(f"❌ שגיאה ביצירת TXT: {str(e)}")

    def generate_docx(self, segments, output_path):
        try:
            doc = Document()
            for seg in segments:
                doc.add_paragraph(seg["text"].strip())
            doc.save(output_path)
            logger.info(f"✅ קובץ DOCX נוצר: {output_path}")
        except Exception as e:
            logger.error(f"❌ שגיאה ביצירת DOCX: {str(e)}")

    def generate_pdf(self, segments, output_path):
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.add_font("Arial", "", fname="C:\\Windows\\Fonts\\arial.ttf", uni=True)
            pdf.set_font("Arial", size=12)
            for seg in segments:
                pdf.multi_cell(0, 10, seg["text"].strip())
            pdf.output(output_path)
            logger.info(f"✅ קובץ PDF נוצר: {output_path}")
        except Exception as e:
            logger.error(f"❌ שגיאה ביצירת PDF: {str(e)}")

    def split_text_smart(self, text, max_chars=80):
        import re
        from textwrap import wrap
        pattern = r'(?<=[.?!،,:;־\-])\s+'
        parts = re.split(pattern, text)
        result = []
        buffer = ""
        for part in parts:
            if len(buffer) + len(part) + 1 <= max_chars:
                buffer += part + " "
            else:
                if buffer:
                    result.append(buffer.strip())
                if len(part) > max_chars:
                    result.extend(wrap(part, max_chars))
                    buffer = ""
                else:
                    buffer = part + " "
        if buffer:
            result.append(buffer.strip())
        return result

    def split_segment(self, seg, max_chars=80):
        chunks = self.split_text_smart(seg["text"], max_chars)
        duration = seg["end"] - seg["start"]
        chunk_duration = duration / len(chunks) if chunks else duration

        new_segments = []
        for i, chunk in enumerate(chunks):
            new_seg = {
                "start": seg["start"] + i * chunk_duration,
                "end": seg["start"] + (i + 1) * chunk_duration,
                "text": chunk.strip()
            }
            new_segments.append(new_seg)
        return new_segments

    def generate_srt(self, segments):
        srt_path = self.config.output_srt
        logger.info(f"📄 יצירת קובץ SRT: {srt_path}")
        try:
            with open(srt_path, "w", encoding="utf-8") as f:
                idx = 1
                for seg in segments:
                    split_segs = self.split_segment(seg) if len(seg["text"]) > 80 else [seg]
                    for part in split_segs:
                        start = self.format_timestamp(part['start'])
                        end = self.format_timestamp(part['end'])
                        text = part['text']
                        f.write(f"{idx}\n{start} --> {end}\n{text}\n\n")
                        idx += 1
            logger.info("✅ קובץ SRT נוצר בהצלחה")
        except Exception as e:
            logger.error(f"❌ שגיאה ביצירת SRT: {str(e)}")
            raise



    def clean_temp_files(self):
        try:
            if os.path.exists(self.config.output_audio):
                os.remove(self.config.output_audio)
                logger.info(f"🧹 נמחק קובץ זמני: {self.config.output_audio}")
        except Exception as e:
            logger.warning(f"⚠️ בעיה במחיקת קובץ זמני: {str(e)}")

    def is_video_file(self, filepath):
        return os.path.splitext(filepath)[1].lower() in [".mp4", ".mov", ".avi", ".webm"]

    def transcribe_and_process(self, filepath, user_id=None):
        # ✅ הוספת בדיקה בתחילת הפונקציה
        if user_id:
            user_plan, current_usage = get_user_usage_and_plan(user_id)
            if not is_allowed_to_transcribe(user_plan, current_usage):
                return {
                    "status": "denied",
                    "message": "חרגת מהמכסה במסלול שלך",
                    "allowed": False
                }
        print("⚙️ נכנסנו ל־transcribe_and_process()")
        logger.info(f"▶️ התחלת תהליך תמלול עבור: {filepath}")

        try:
            update_status("extracting_audio", 10)
            print("🎧 extracting audio...")
            self.extract_audio()

            # חישוב משך
            try:
                self.duration = self.get_video_duration(filepath)
                print("⏱️ video duration =", self.duration)
            except Exception as e:
                logger.warning(f"⚠️ שגיאה בחישוב משך וידאו: {e}")
                self.duration = 0

            update_status("transcribing", 30)
            print("📝 מתחיל תמלול...")
            try:
                segments = self.transcribe_audio()
                print(f"📚 תמלול הניב {len(segments)} סגמנטים")
            except Exception as e:
                logger.error(f"❌ שגיאה בתמלול: {str(e)}")
                import traceback
                print(traceback.format_exc())
                return {
                    "status": "failed",
                    "message": str(e),
                    "traceback": traceback.format_exc(),
                    "duration": getattr(self, 'duration', 0)
                }

            if not segments:
                logger.error("❌ לא התקבלו סגמנטים מהתמלול")
                return {
                    "status": "failed",
                    "message": "No transcription segments returned",
                    "duration": self.duration
                }

            self.detected_lang = getattr(self, 'source_language', "unknown")

            # יצירת שמות קבצים
            base_name = os.path.splitext(os.path.basename(self.config.output_video))[0]
            task_id = base_name.split("_")[-1]
            txt_path = os.path.join("results", f"{base_name}.txt")
            docx_path = os.path.join("results", f"{base_name}.docx")
            pdf_path = os.path.join("results", f"{base_name}.pdf")
            srt_path = self.config.output_srt
            video_path = self.config.output_video

            # יצירת קבצים
            update_status("creating_documents", 50)
            print("🛠️ מייצר txt/docx/pdf...")
            self.generate_txt(segments, txt_path)
            self.generate_docx(segments, docx_path)
            self.generate_pdf(segments, pdf_path)

            update_status("generating_srt", 75)
            print("🎬 מייצר SRT...")
            self.generate_srt(segments)

            # העלאה ל־R2
            update_status("uploading", 95)
            print("🔼 מעלה קבצי תוצאה ל־R2...")
            print("🧪 debug output_video =", self.config.output_video)
            print("🧪 base_name =", base_name)
            print("🧪 task_id =", task_id)
            print("🧪 קבצים בתיקיית results:", os.listdir("results"))

            R2_PUBLIC_BASE = os.getenv("R2_PUBLIC_BASE") or \
                             "https://eb0e988de4e9fb026f45d9e3038314e2.r2.cloudflarestorage.com/talkscribe-uploads"
            print("📂 תוכן תיקיית results:", os.listdir("results"))

            def try_upload(path, r2_key, label):
                if os.path.exists(path):
                    success = upload_file_to_r2(path, r2_key)
                    print(f"✅ {label} uploaded:", success)
                else:
                    print(f"❌ {label} לא קיים:", path)

            try_upload(txt_path, f"{task_id}/{base_name}.txt", "TXT")
            try_upload(docx_path, f"{task_id}/{base_name}.docx", "DOCX")
            try_upload(pdf_path, f"{task_id}/{base_name}.pdf", "PDF")
            try_upload(srt_path, f"{task_id}/{base_name}.srt", "SRT")

            if self.is_video_file(filepath):
                try:
                    print("🎞️ צורב כתוביות לסרטון...")
                    ass_path = srt_path.replace(".srt", ".ass")
                    DEFAULT_SUBTITLE_STYLE = {
                        "font": "Arial",
                        "size": 36,
                        "alignment": "bottom_center",
                        "primary_color": "#FFFFFF",
                        "outline_color": "#000000",
                        "outline": 2,
                        "shadow": 0,
                        "margin_v": 60
                    }

                    style_cfg = getattr(self.config, "subtitle_style", None) or DEFAULT_SUBTITLE_STYLE
                    rtl = self.detected_lang in ["he", "ar", "fa"]
                    convert_srt_to_ass(srt_path, ass_path, style_config=style_cfg, rtl=rtl)
                    burn_ass_subtitles(filepath, ass_path, video_path)
                    try_upload(video_path, f"{task_id}/{base_name}.mp4", "MP4")
                except Exception as e:
                    logger.error(f"❌ שגיאה בצריבת כתוביות: {str(e)}")
                    import traceback
                    print(traceback.format_exc())

            # ניקוי
            update_status("cleaning", 98)
            self.clean_temp_files()

            update_status("done", 100)
            logger.info(f"🎉 סיום: {len(segments)} סגמנטים, משך {self.duration} שניות")

            # ✍️ עדכון שימוש למשתמש
            if user_id and self.duration > 0:

                default_subtitle_style = {
                    "font": "Arial",
                    "font_size": 26,
                    "primary_color": "#FFFFFF",
                    "outline_color": "#000000",
                    "alignment": 2,
                    "background_color": "transparent"
                }

                transcriptions_collection.insert_one({
                    "user_id": user_id,
                    "task_id": task_id,
                    "file_name": os.path.basename(filepath),
                    "duration": self.duration,
                    "language": self.detected_lang,
                    "transcript_text": " ".join([s["text"] for s in segments]),
                    "created_at": datetime.utcnow(),
                    "subtitle_style": default_subtitle_style,
                    "r2_urls": {
                        "txt": f"{R2_PUBLIC_BASE}/{task_id}/{base_name}.txt",
                        "docx": f"{R2_PUBLIC_BASE}/{task_id}/{base_name}.docx",
                        "pdf": f"{R2_PUBLIC_BASE}/{task_id}/{base_name}.pdf",
                        "srt": f"{R2_PUBLIC_BASE}/{task_id}/{base_name}.srt",
                        "mp4": f"{R2_PUBLIC_BASE}/{task_id}/{base_name}.mp4" if self.is_video_file(filepath) else None
                    }
                })
                logger.info(f"🔄 מעדכן שימוש למשתמש: {user_id}, דקות: {self.duration}")
                print(f"📣 Calling update_user_usage... (duration={self.duration})")
                success = update_user_usage(user_id, self.duration)
                print("✅ update_user_usage success:", success)
            else:
                print(f"⚠️ לא עודכן שימוש — user_id={user_id}, duration={self.duration}")

            return {
                "status": "completed",
                "task_id": task_id,
                "base_name": base_name,
                "duration": self.duration,
                "detected_language": self.detected_lang,
                "transcript_text": " ".join([s["text"] for s in segments]),
                "r2_urls": {
                    "txt": f"{R2_PUBLIC_BASE}/{task_id}/{base_name}.txt",
                    "docx": f"{R2_PUBLIC_BASE}/{task_id}/{base_name}.docx",
                    "pdf": f"{R2_PUBLIC_BASE}/{task_id}/{base_name}.pdf",
                    "srt": f"{R2_PUBLIC_BASE}/{task_id}/{base_name}.srt",
                    "mp4": f"{R2_PUBLIC_BASE}/{task_id}/{base_name}.mp4" if self.is_video_file(filepath) else None
                }
            }

        except Exception as e:
            logger.error(f"❌ שגיאה בתהליך כולו: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return {
                "status": "failed",
                "message": str(e),
                "traceback": traceback.format_exc(),
                "duration": getattr(self, 'duration', 0)
            }

    import subprocess
    import json



    def get_video_duration(self, path):
        try:
            result = subprocess.run(
                [self.FFPROBE_PATH, "-v", "error", "-show_entries", "format=duration",
                 "-of", "json", path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True
            )
            data = json.loads(result.stdout)
            duration = float(data["format"]["duration"])
            return int(duration)
        except Exception as e:
            print(f"⚠️ ffprobe נכשל: {e}")
            return 0

def get_video_resolution(video_path):
    cmd = [
        "ffprobe",
        "-v", "error",
        "-select_streams", "v:0",
        "-show_entries", "stream=width,height",
        "-of", "csv=s=x:p=0",
        video_path
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode == 0:
        res = result.stdout.strip()
        width, height = map(int, res.split('x'))
        return (width, height)
    else:
        return None

def generate_outputs_from_srt(srt_path: str, task_id: str):
    from docx import Document
    from fpdf import FPDF
    from pathlib import Path

    base_name = task_id
    output_dir = Path("results")
    output_dir.mkdir(parents=True, exist_ok=True)

    # קרא את ה־SRT
    with open(srt_path, "r", encoding="utf-8") as f:
        content = f.read()

    # שלב 1: חילוץ הטקסט (בלי מספרים וזמנים)
    import re
    blocks = re.split(r"\n\s*\n", content)
    lines = []
    for block in blocks:
        parts = block.strip().splitlines()
        if len(parts) >= 3:
            lines.append(" ".join(parts[2:]))  # שורת הטקסט
        elif len(parts) == 2:
            lines.append(parts[1])
    full_text = "\n".join(lines).strip()

    # יצירת קובץ TXT
    txt_path = output_dir / f"{base_name}.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(full_text)

    # יצירת קובץ DOCX
    docx_path = output_dir / f"{base_name}.docx"
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(docx_path)

    # יצירת קובץ PDF
    pdf_path = output_dir / f"{base_name}.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("Arial", "", fname="C:\\Windows\\Fonts\\arial.ttf", uni=True)
    pdf.set_font("Arial", size=12)
    for line in lines:
        pdf.multi_cell(0, 10, line)
    pdf.output(str(pdf_path))

    return [str(txt_path), str(docx_path), str(pdf_path)]

# 📦 צורב כתוביות מה־MongoDB לפי task_id
def embed_subtitles_from_db(task_id):
    doc = transcriptions_collection.find_one({"task_id": task_id})
    if not doc:
        raise ValueError(f"לא נמצא תמלול עם task_id: {task_id}")

    style_cfg = doc.get("subtitle_style", {})
    srt_path = f"results/{task_id}.srt"
    ass_path = srt_path.replace(".srt", ".ass")
    input_video = f"uploads/{task_id}.mp4"  # ודא שזה הנתיב בו אתה שומר את המקור
    output_video = f"results/{task_id}_styled.mp4"

    rtl = doc.get("language") in ["he", "ar", "fa"]

    convert_srt_to_ass(srt_path, ass_path, style_config=style_cfg, rtl=rtl)
    burn_ass_subtitles(input_video, ass_path, output_video)

    return output_video
